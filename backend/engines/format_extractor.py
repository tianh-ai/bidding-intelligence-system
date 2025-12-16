"""
文档格式信息提取器
功能：
1. 提取DOCX文件的字体信息（字号、字体名称、加粗、斜体）
2. 提取段落格式（对齐方式、行距、段落间距、缩进）
3. 提取页面布局信息（页边距、纸张大小）
4. 保存到structure_data字段
"""

import docx
from typing import Dict, List, Optional, Any
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FormatExtractor:
    """格式信息提取器"""
    
    def __init__(self):
        """初始化提取器"""
        pass
    
    def extract_format_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        从DOCX文件提取格式信息
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            Dict包含文档级别的格式统计和段落格式列表
        """
        try:
            doc = docx.Document(file_path)
            
            # 提取文档级别信息
            doc_format = {
                'page_setup': self._extract_page_setup(doc),
                'styles_used': self._extract_styles_used(doc),
                'paragraphs': [],
                'font_statistics': {},  # 字体统计
            }
            
            # 提取每个段落的格式
            font_counts = {}
            for para_idx, para in enumerate(doc.paragraphs):
                para_format = self._extract_paragraph_format(para, para_idx)
                if para_format['content']:  # 只记录有内容的段落
                    doc_format['paragraphs'].append(para_format)
                    
                    # 统计字体使用
                    font_name = para_format.get('font', {}).get('name', '未知')
                    font_counts[font_name] = font_counts.get(font_name, 0) + 1
            
            # 生成字体统计
            doc_format['font_statistics'] = {
                'most_common_font': max(font_counts, key=font_counts.get) if font_counts else None,
                'font_usage': font_counts,
                'total_paragraphs': len(doc_format['paragraphs']),
            }
            
            return doc_format
            
        except Exception as e:
            from core.logger import logger
            logger.error(f"格式提取失败: {e}")
            return {'error': str(e)}
    
    def _extract_page_setup(self, doc: docx.Document) -> Dict[str, Any]:
        """提取页面设置"""
        try:
            section = doc.sections[0] if doc.sections else None
            if not section:
                return {}
            
            return {
                'page_width': self._convert_to_cm(section.page_width),
                'page_height': self._convert_to_cm(section.page_height),
                'margin_top': self._convert_to_cm(section.top_margin),
                'margin_bottom': self._convert_to_cm(section.bottom_margin),
                'margin_left': self._convert_to_cm(section.left_margin),
                'margin_right': self._convert_to_cm(section.right_margin),
                'orientation': 'landscape' if section.page_width > section.page_height else 'portrait',
            }
        except Exception as e:
            return {'error': str(e)}
    
    def _extract_styles_used(self, doc: docx.Document) -> List[str]:
        """提取文档中使用的样式列表"""
        styles = set()
        for para in doc.paragraphs:
            if para.style and para.style.name:
                styles.add(para.style.name)
        return sorted(list(styles))
    
    def _extract_paragraph_format(self, para: docx.text.paragraph.Paragraph, index: int) -> Dict[str, Any]:
        """提取段落格式"""
        try:
            # 提取段落文本
            text_content = para.text.strip()
            
            # 提取段落格式
            para_fmt = para.paragraph_format
            
            # 提取字体格式（取第一个run的格式作为代表）
            font_info = {}
            if para.runs:
                first_run = para.runs[0]
                font = first_run.font
                font_info = {
                    'name': font.name or '默认',
                    'size': self._convert_font_size(font.size),
                    'bold': font.bold,
                    'italic': font.italic,
                    'underline': font.underline,
                    'color': self._extract_color(font.color),
                }
            
            return {
                'index': index,
                'content': text_content,
                'content_length': len(text_content),
                'style': para.style.name if para.style else None,
                'font': font_info,
                'alignment': self._get_alignment_name(para_fmt.alignment),
                'line_spacing': self._get_line_spacing(para_fmt.line_spacing),
                'space_before': self._convert_to_pt(para_fmt.space_before),
                'space_after': self._convert_to_pt(para_fmt.space_after),
                'left_indent': self._convert_to_pt(para_fmt.left_indent),
                'right_indent': self._convert_to_pt(para_fmt.right_indent),
                'first_line_indent': self._convert_to_pt(para_fmt.first_line_indent),
            }
            
        except Exception as e:
            return {
                'index': index,
                'content': para.text.strip() if para else '',
                'error': str(e),
            }
    
    def _convert_to_cm(self, value) -> Optional[float]:
        """将EMU单位转换为厘米"""
        if value is None:
            return None
        try:
            # EMU to Inches to CM
            inches = value / 914400  # 1 inch = 914400 EMU
            cm = inches * 2.54
            return round(cm, 2)
        except:
            return None
    
    def _convert_to_pt(self, value) -> Optional[float]:
        """将EMU单位转换为磅（pt）"""
        if value is None:
            return None
        try:
            # EMU to Points
            pt = value / 12700  # 1 pt = 12700 EMU
            return round(pt, 1)
        except:
            return None
    
    def _convert_font_size(self, size) -> Optional[float]:
        """转换字体大小为磅"""
        if size is None:
            return None
        try:
            # size是Pt对象
            return round(size.pt, 1) if hasattr(size, 'pt') else float(size) / 12700
        except:
            return None
    
    def _get_alignment_name(self, alignment) -> str:
        """获取对齐方式名称"""
        if alignment is None:
            return 'left'
        
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: 'distribute',
        }
        
        return alignment_map.get(alignment, 'left')
    
    def _get_line_spacing(self, line_spacing) -> Optional[float]:
        """获取行距"""
        if line_spacing is None:
            return None
        try:
            # line_spacing 是 Length 对象
            return round(line_spacing, 2) if isinstance(line_spacing, (int, float)) else None
        except:
            return None
    
    def _extract_color(self, color) -> Optional[str]:
        """提取颜色"""
        try:
            if color and color.rgb:
                return f"#{color.rgb}"
            return None
        except:
            return None
    
    def extract_chapter_formats(
        self,
        file_path: str,
        chapters: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        为每个章节提取格式信息
        
        Args:
            file_path: DOCX文件路径
            chapters: 章节列表（包含start_line和end_line）
            
        Returns:
            每个章节的格式信息列表
        """
        try:
            doc = docx.Document(file_path)
            doc_format = self.extract_format_from_docx(file_path)
            
            # 为每个章节分配对应的段落格式
            chapter_formats = []
            for chapter in chapters:
                start_line = chapter.get('start_line', 0)
                end_line = chapter.get('end_line', len(doc.paragraphs))
                
                # 提取该章节范围内的段落格式
                chapter_paras = [
                    p for p in doc_format.get('paragraphs', [])
                    if start_line <= p.get('index', 0) < end_line
                ]
                
                # 统计该章节的格式特征
                chapter_format = {
                    'chapter_number': chapter.get('chapter_number'),
                    'chapter_title': chapter.get('chapter_title'),
                    'paragraph_count': len(chapter_paras),
                    'paragraphs': chapter_paras,
                    'dominant_font': self._get_dominant_font(chapter_paras),
                    'dominant_alignment': self._get_dominant_alignment(chapter_paras),
                }
                
                chapter_formats.append(chapter_format)
            
            return chapter_formats
            
        except Exception as e:
            from core.logger import logger
            logger.error(f"章节格式提取失败: {e}")
            return []
    
    def _get_dominant_font(self, paragraphs: List[Dict]) -> Optional[Dict]:
        """获取主要使用的字体"""
        font_counts = {}
        for para in paragraphs:
            font_info = para.get('font', {})
            font_name = font_info.get('name')
            font_size = font_info.get('size')
            
            if font_name and font_size:
                key = f"{font_name}_{font_size}"
                font_counts[key] = font_counts.get(key, 0) + 1
        
        if not font_counts:
            return None
        
        dominant_key = max(font_counts, key=font_counts.get)
        font_name, font_size = dominant_key.rsplit('_', 1)
        
        return {
            'name': font_name,
            'size': float(font_size),
            'usage_count': font_counts[dominant_key],
        }
    
    def _get_dominant_alignment(self, paragraphs: List[Dict]) -> str:
        """获取主要使用的对齐方式"""
        alignment_counts = {}
        for para in paragraphs:
            alignment = para.get('alignment', 'left')
            alignment_counts[alignment] = alignment_counts.get(alignment, 0) + 1
        
        if not alignment_counts:
            return 'left'
        
        return max(alignment_counts, key=alignment_counts.get)


# 单例获取函数
_format_extractor_instance = None

def get_format_extractor() -> FormatExtractor:
    """获取格式提取器实例"""
    global _format_extractor_instance
    if _format_extractor_instance is None:
        _format_extractor_instance = FormatExtractor()
    return _format_extractor_instance
