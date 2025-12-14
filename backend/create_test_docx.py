"""
创建一个包含图片的测试DOCX文件并测试图片提取
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent / "backend"))

from docx import Document
from docx.shared import Inches
from PIL import Image
import uuid
import io
from datetime import datetime


def create_test_docx_with_images():
    """创建包含3张图片的测试DOCX"""
    
    print("创建测试DOCX文件...")
    
    # 1. 创建文档
    doc = Document()
    doc.add_heading('测试文档 - 包含图片', 0)
    
    doc.add_paragraph('这是一个包含多张图片的测试文档，用于验证图片提取功能。')
    
    # 2. 创建并添加3张测试图片
    for i in range(1, 4):
        # 创建一张简单的彩色图片
        img = Image.new('RGB', (400, 300), color=(255 * (i % 2), 128, 255 - 50 * i))
        
        # 添加文本到图片
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 48)
        except:
            font = ImageFont.load_default()
        
        draw.text((100, 120), f"Test Image {i}", fill=(255, 255, 255), font=font)
        
        # 保存到内存
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        # 添加到文档
        doc.add_heading(f'图片 {i}', level=2)
        doc.add_picture(img_bytes, width=Inches(3))
        doc.add_paragraph(f'这是第{i}张测试图片。')
    
    # 3. 保存文档
    test_dir = Path("/Volumes/ssd/bidding-data/uploads/temp")
    test_dir.mkdir(parents=True, exist_ok=True)
    
    test_file = test_dir / f"test_with_images_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(test_file)
    
    print(f"✅ 测试文件已创建: {test_file}")
    print(f"   文件大小: {test_file.stat().st_size} bytes")
    
    return str(test_file)


def test_extraction(docx_path):
    """测试图片提取"""
    from engines.image_extractor import ImageExtractor
    from database import db
    
    print(f"\n开始测试图片提取...")
    print(f"文件: {Path(docx_path).name}")
    
    extractor = ImageExtractor()
    file_id = str(uuid.uuid4())
    year = 2025
    
    # 提取图片
    images = extractor.extract_from_docx(docx_path, file_id, year)
    
    print(f"\n✅ 提取了 {len(images)} 张图片\n")
    
    for img in images:
        print(f"图片 {img['image_number']}:")
        print(f"  格式: {img['format']}")
        print(f"  尺寸: {img['width']}x{img['height']}")
        print(f"  大小: {img['size']} bytes")
        print(f"  路径: {img['image_path']}")
        print(f"  Hash: {img['hash']}")
        
        # 验证文件存在
        if Path(img['image_path']).exists():
            print(f"  ✅ 文件已保存到磁盘")
        else:
            print(f"  ❌ 文件不存在")
        print()
    
    # 验证数据库
    db_count = db.query_one(
        "SELECT COUNT(*) as count FROM extracted_images WHERE file_id = %s",
        (file_id,)
    )
    
    print(f"数据库记录数: {db_count['count']}")
    
    if db_count['count'] == len(images):
        print("✅ 数据库记录正确")
    else:
        print(f"❌ 数据库记录不匹配 (期望{len(images)}, 实际{db_count['count']})")
    
    return images


if __name__ == "__main__":
    try:
        docx_path = create_test_docx_with_images()
        images = test_extraction(docx_path)
        
        print(f"\n{'='*60}")
        print(f"测试完成! 成功提取 {len(images)} 张图片")
        print(f"{'='*60}")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
